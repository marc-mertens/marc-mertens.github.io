# -*- coding: utf-8 -*-
"""
Created on Mon Jul  5 15:54:45 2021

@author: mertensm
docx: https://automatetheboringstuff.com/chapter13/
fuzzy string matching: https://www.datacamp.com/community/tutorials/fuzzy-string-python
"""

import docx
import json
#doc = docx.Document('Recepten-2021.docx')
#print(len(doc.paragraphs))
# doc.paragraphs[0].text

tags=["Aantal personen", "Aantal stuks","Ingrediënten", "Bereiding", "Afwerking", "Bron",
      "Baktijd","Temperatuur","Voorbereiding","Opdienen","Tips","Kooktijd","Tip"]

# read whole document in a string
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

#print(getText('Recepten-2021.docx'))
    
# read all headers in a string
def Docx2Json(filename):
    doc = docx.Document(filename)
    data = []
    i=-1
    context=''
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
           category = para.text
        elif para.style.name == 'Heading 2':
           data.append({})
           i=i+1
           data[i]['category'] = category
           data[i]['name'] = para.text
        elif para.text.startswith('Aantal personen:'):
           data[i]['portions'] = para.text.replace('Aantal personen:','').strip()
        elif para.text.startswith('Ingrediënten:'):
           data[i]['ingredients'] = para.text.replace('Ingrediënten:','').strip()
        elif para.text.startswith('Bereiding:'):
           data[i]['preparation'] = para.text.replace('Bereiding:','').strip()   
    return json.dumps(data,indent=2)

def GetTag(atext):
    atag = atext[0:50][0:atext.find(':')]
    return atag

def DocxListTags(filename):
    doc = docx.Document(filename)
    data = []
    recipe_num = -1
    context = ''
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
           category = para.text
        elif para.style.name == 'Heading 2':
           title = para.text
           context = 'started'
        elif context == 'started' and not (GetTag(para.text) in tags):
           data.append({})
           recipe_num = recipe_num + 1
           data[recipe_num]['category'] = category
           data[recipe_num]['name'] = title   
           data[recipe_num]['tag'] = GetTag(para.text)
    return json.dumps(data,indent=2)

def GetWithoutTag(atext):
    therest = atext[atext.find(':')+1:].strip()
    return therest

def RemoveEscape(atext):
    return atext.replace('\n','<br>').replace('\r','').replace('\t',' ').replace('\"','\u00B4').replace("\'","\u00B4")
    
def Recepten2Json(receptfilename, jsonfilename):
    doc = docx.Document(receptfilename)
    data = []  
    context = ''     #avoids that text before the first Heading 2 would be processed
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':  # start of a new chapter (=group with the same category)
           category = RemoveEscape(para.text)
           context=''
        elif para.style.name == 'Heading 2': # start of a new recipe
           data.append({})
           data[-1]['category'] = category
           data[-1]['name'] = RemoveEscape(para.text)   
           context = 'in recipe'
           error_count = 0 #num unprocessed blocks
        elif context != '':  # there is data to be processed
           atag = GetTag(para.text)
           if (atag in tags): # start of a new section, starting with a tag
               context=atag
               if context == "Ingrediënten" :
                   if para.text.count(',') > para.text.count(';'):
                       delim = ','
                   else:
                       delim = ';'
                   data[-1][context] = [x.strip() for x in RemoveEscape(GetWithoutTag(para.text)).replace(':',':' + delim).split(delim)]                  
               else: # it is a simple text block
                   data[-1][context] = RemoveEscape(GetWithoutTag(para.text))
           else: #no tag, so this is a part of the previous tag
               if context == 'Ingrediënten' :
                   data[-1][context].extend([x.strip() for x in RemoveEscape(GetWithoutTag(para.text)).replace(':',':' + delim).split(delim)])
               elif context != 'in recipe' : # in tags :    
                   data[-1][context] += '<br>' + RemoveEscape(para.text)
               else: # error ????
                   error_count +=1
                   data[-1]['error_'+ str(error_count)] = RemoveEscape(para.text)
    jsonfile = open(jsonfilename,"w")    
    jsonstring = json.dumps(sorted(data, key=lambda x: x['name'].lower())) #,indent=2)
    jsonfile.writelines("data='" + jsonstring + "';")
    jsonfile.close()
    return str(len(data)) + ' recipes processed'

print(Recepten2Json('..\Recepten-2021.docx','data.js'))

#print(Docx2Json('Recepten-2021.docx'))
#file1 = open("ReceptenTest.txt","w")    
#print(DocxListTags('Recepten-2021.docx'))
#file1.writelines(DocxListTags('Recepten-2021.docx'))
#file1.close()